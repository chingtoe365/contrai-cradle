# Contract preprocessing script
# Author: Jintao Long <jintaolong@brookes.ac.uk>

__doc__ = """
	This script is used for preprocessing contracts 
	with the following stepstones
	- tokenization
	- stop words removing
	- stemming
	- speech tagging
	- trunking
	- chinking
	- named entity recognition
	- lemmatizing
"""

import nltk
nltk.download('stopwords')
import docx
import os
import abc
import re, string, unicodedata
import datetime
import json
import pickle
import contractions
import inflect
import pandas as pd
from bs4 import BeautifulSoup


from typing import Callable, Dict, List
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from sklearn.feature_extraction.text import TfidfVectorizer

from utils import clean_string, update_dict_by_accumulate
from word_embedding import *
from tagging import *
import logging
from event_logger import logger

ps = PorterStemmer()
TRAINING_INGREDIENT_PATH = 'training_ingredients/'
MIDDLE_INGREDIENT_PATH = 'middleware_ingredients/'
WORD2VEC_INGREDIENT_PATH = 'word2vec_ingredients/'

class PreprocessingAbstract():
	"""
	Bundle of preprocessing methods
	"""
	def __init__(self, file_path, embedding_method, tag_obtaining_method,
			remove_stop_words: bool, do_stemming: bool, strigent_topic: bool,
			ngram: int, multiple_paragraphs: bool, ngram_mixed:bool):
		self._file_path = file_path
		self._contract_doc = self._load_file(self._file_path)
		self._remove_stop_words = remove_stop_words
		self._do_stemming = do_stemming
		self._stop_words = set(stopwords.words("english"))
		self._embedding_method = embedding_method
		self._tag_obtaining_method = eval(tag_obtaining_method)
		self._strigent_topic = strigent_topic
		self._ngram = ngram
		self._multiple_paragraphs = multiple_paragraphs
		self._ngram_mixed = ngram_mixed

	@abc.abstractmethod
	def _load_file(self, file_path):
		"""
		Load file
		"""
		pass
	
	def _strip_html(self, text):
	    soup = BeautifulSoup(text, "html.parser")
	    return soup.get_text()

	def _remove_between_square_brackets(self, text):
	    return re.sub(r'\[[^]]*\]', '', text)

	def denoise_text(self, text):
	    text = self._strip_html(text)
	    text = self._remove_between_square_brackets(text)
	    return text
	
	def replace_contractions(self, text):
	    """Replace contractions in string of text"""
	    return contractions.fix(text)
	
	def _remove_non_ascii(self, words):
	    """Remove non-ASCII characters from list of tokenized words"""
	    new_words = []
	    for word in words:
	        new_word = unicodedata.normalize('NFKD', word).encode(
	        	'ascii', 'ignore').decode('utf-8', 'ignore')
	        new_words.append(new_word)
	    return new_words

	def _to_lowercase(self, words):
	    """Convert all characters to lowercase from list of tokenized words"""
	    new_words = []
	    for word in words:
	        new_word = word.lower()
	        new_words.append(new_word)
	    return new_words

	def _remove_punctuation(self, words):
	    """Remove punctuation from list of tokenized words"""
	    new_words = []
	    for word in words:
	        new_word = re.sub(r'[^\w\s]', '', word)
	        if new_word != '':
	            new_words.append(new_word)
	    return new_words

	def _replace_numbers(self, words):
	    """Replace all interger occurrences in list of tokenized words with textual representation"""
	    p = inflect.engine()
	    new_words = []
	    for word in words:
	        if word.isdigit():
	            new_word = p.number_to_words(word)
	            new_words.append(new_word)
	        else:
	            new_words.append(word)
	    return new_words

	def normalize(self, words):
	    words = self._remove_non_ascii(words)
	    words = self._to_lowercase(words)
	    words = self._remove_punctuation(words)
	    words = self._replace_numbers(words)
	    return words

	def _tokenize(self, text):
		"""
		Tokenize the entire contract by sentences 
		"""
		return nltk.word_tokenize(text)

	def _is_stop_word(self, word):
		"""
		Remove stop words like for/in/at/... etc.
		Because they usually have no meanings

		TODO: more specific rules for legal contract stop words
		"""
		return word.lower() in self._stop_words

	def _stemming(self, word):
		"""
		Do stemming to make for example,
		doing==do, better==good
		because they have usually same meanings

		TODO: more specific stemming rules for legal contract
		"""
		return ps.stem(word)

	def _word_embedding(
			self, method: Callable, list_of_words: List) -> Dict:
		"""
		word embedding that turn words into quantitative values 
		which can then be machine learned
		"""
		return method(list_of_words)


	def _get_tags(
			self, quant_word_vec: Dict, word_vec: Dict, 
			current_label: str, method: Callable, 
			strigent: bool) -> Dict:
		"""
		Get tags for topic classification 
		"""
		return method(quant_word_vec, word_vec, current_label, strigent)

	def _definition_check(self, text):
		return any([x in text for x in ["definition", "intepretation"]])

	def _non_legal_content_check(self, paragraph):
		"""
		To identify paragraphs with non-legal content

		Starting with words like: 
		#contents #appendix #annex #exhibit #schedule
		"""
		return 	re.search(r'^contents', paragraph) or \
			re.search(r'^appendix', paragraph) or \
			re.search(r'^annex', paragraph) or \
			re.search(r'^exhibit', paragraph) or \
			re.search(r'^schedule', paragraph) or  \
			re.search(r'^list of', paragraph) or \
			re.search(r'^signed by', paragraph) or \
			re.search(r'^service block statement of work', paragraph)

	def _caption_as_label(
			self,
			raw_text: str, 
			word_vec: List[str], 
			label_wc_limit=7, 
			stop_word_count_limit=2, 
		) -> bool:
		# search up-level caption as tag/topic
		"""
		Rules:
		- stop words not more than 2/3, excluding AND/OF/THE/FOR
		- no colons
		- paragraph length not more than 6

		TODO: DEPRECIATE when manual tagging is ready
		"""
		is_label = False
		stop_word_count = 0
		if len(word_vec) >= label_wc_limit:
			return is_label
		for word in word_vec:
			if word == ':':
				return is_label
			if word.lower() in ('and', 'of', 'the', 'for'):
				continue
			if self._is_stop_word(word):
				stop_word_count += 1
		if stop_word_count <= stop_word_count_limit:
			is_label = True
		return is_label

	@abc.abstractmethod
	def _numerize_texts(self) -> List:
		pass

	def _word_scan(self, raw_word_vec, current_label, strigent) -> List:
		list_of_words = []
		# args_embedding_input = {}
		# switches to remove stop words and do stemming
		if self._remove_stop_words and self._do_stemming:
			for word in raw_word_vec:
				if self._is_stop_word(word):
					continue
				else:
					list_of_words.append(self._stemming(word))
		elif self._remove_stop_words and not self._do_stemming:
			for word in raw_word_vec:
				if self._is_stop_word(word):
					continue
				else:
					list_of_words.append(word)
		elif not self._remove_stop_words and self._do_stemming:
			for word in raw_word_vec:
				list_of_words.append(self._stemming(word))
		else:
			for word in raw_word_vec:
				list_of_words.append(word)

		# if self._embedding_method != 'tfidf':
		# 	args_embedding_input['list_of_words'] = list_of_words
		list_of_words_original_copy = list_of_words.copy()
		if self._ngram_mixed:
			if self._ngram > 1:
				# if more ngram split is required
				for i in range(2,(self._ngram+1)):
					list_of_words_copy = list_of_words_original_copy.copy()
					while len(list_of_words_copy) >= i:
						list_of_words.append(
							' '.join(list_of_words_copy[0:i])
							# list_of_words_copy[0]+' '+list_of_words_copy[1]
						)
						list_of_words_copy = list_of_words_copy[1:]
		else:
			list_of_words_copy = list_of_words_original_copy.copy()
			list_of_words = []
			while len(list_of_words_copy) >= self._ngram:
				list_of_words.append(
					' '.join(list_of_words_copy[0:i])
					# list_of_words_copy[0]+' '+list_of_words_copy[1]
				)
				list_of_words_copy = list_of_words_copy[1:]
		method = eval(self._embedding_method)
		quantified_word_vec = self._word_embedding(
			method, list_of_words
		)

		# add tag for the paragraph
		return self._get_tags(
			quantified_word_vec, 
			raw_word_vec, 
			current_label,
			self._tag_obtaining_method,
			strigent
		)

	def bag_of_word_dict_transformer(self) -> List:
		"""
		Main function to transform words in contracts 
		to machine learnable format of input. 

		derived observation looks like this:
		[{
			{
				# clause/document 1
				
				"hello" : 1,
				"world": 5,
				"you": 4,
				xxx
			},
			"Indemity"
		}, {
			{
				# clause/document 2

				"legal": 4,
				"fine": 2,
				"responsibility": 1,
				xxx	
			},
			"Termination"
		}, 
		xxxx
		]
		"""

		numeric_objects = self._numerize_texts()
		# save the list
		if 'derived_observations' in numeric_objects:
			self._write_json_to_file(
				TRAINING_INGREDIENT_PATH, 
				numeric_objects['derived_observations']
			)
		if 'topic_file_pointer' in numeric_objects:
			self._write_json_to_file(
				MIDDLE_INGREDIENT_PATH, 
				numeric_objects['topic_file_pointer']
			)
		if 'word2vec_corpus' in numeric_objects:
			self._write_list_to_file(
				WORD2VEC_INGREDIENT_PATH, 
				numeric_objects['word2vec_corpus']
			)


	def _write_json_to_file(self, path_to_write, output: List[Dict]):
		"""
		Save list of dictionaries and write to txt file
		"""
		filebase = os.path.splitext(os.path.basename(
				self._file_path))[0]
		filename = filebase
		filepath = os.path.join(
			os.getcwd(), 
			path_to_write, 
			filebase+'.txt'
		)
		f = open(filepath, 'w')
		f.write(json.dumps(output))
		f.close()

	def _write_list_to_file(self, path_to_write, output: List[List[str]]):
		"""
		Save list of dictionaries and write to txt file
		"""
		filebase = os.path.splitext(os.path.basename(
				self._file_path))[0]
		filename = filebase
		filepath = os.path.join(
			os.getcwd(), 
			path_to_write, 
			filebase+'.data'
		)
		f = open(filepath, 'wb')
		pickle.dump(output, f)
		f.close()


class DocxPreprocessing(PreprocessingAbstract):
	def _load_file(self, file_path):
		full_path = os.path.join(os.getcwd(), file_path)
		return docx.Document(full_path)


	def _numerize_texts(self) -> List:
		derived_observations = []
		word2vec_corpus = []
		topic_file_pointer = []
		current_label = ''
		if len(self._contract_doc.paragraphs) > 0:
			for p in self._contract_doc.paragraphs:
				word_vec = []
				if len(p.text) == 0:
					# if the paragraph is empty skip it
					continue

				# if a paragraph starts with Appendix then jump out 
				# and stop looping through the rest of doc as those are not needed
				# trait: page title

				if p.text.lower().startswith('appendix') or \
				p.text.lower().startswith('exhibit') or \
				p.text.lower().startswith('annex') or \
				p.text.lower().startswith('schedule'):
					break

				# denoise text
				text = self.denoise_text(p.text)
				# replace contractions
				# eg. don't --> do n't
				text = self.replace_contractions(text)
				# remove curly brackets
				text = re.sub(r'[{,}]', '', text)
				raw_word_vec = self._tokenize(text)
				# normalize the tokens
				raw_word_vec = self.normalize(raw_word_vec)
				# make TAG_OR_NOT judgement first as we use stop words to identify
				# TODO: REMOVE this when we have mannual tags in documents
				if self._caption_as_label(text, raw_word_vec):
					current_label = clean_string(text)

				derived_observations.append(
					self._word_scan(
						raw_word_vec, current_label, self._strigent_topic
					)
				)								
				topic_file_pointer.append(
					self._word_scan(
						raw_word_vec, current_label, self._strigent_topic
					) + \
					(
						current_label,
						os.path.basename(self._file_path),
					)
				)
				if self._do_stemming:
					for word in raw_word_vec:
						word_vec.append(self._stemming(word))
				else:
					word_vec = raw_word_vec
				word2vec_corpus.append(word_vec)

		return {
			'derived_observations': derived_observations,
			'topic_file_pointer': topic_file_pointer, 
			'word2vec_corpus': word2vec_corpus
		}


class RtfPreprocessing(PreprocessingAbstract):
	_paragraph_separator_sub = "###############"
	def _load_file(self, file_path):
		full_path = os.path.join(os.getcwd(), file_path)
		return open(full_path).read()

	def _numerize_texts(self) -> List:
		derived_observations = []
		same_topic_BOW = []
		# add a file to point topic to file
		# only used as reference for research
		topic_file_pointer = []
		tfidf_calculation_corpus = []
		tfidf_calculation_topic = []
		current_label = ''
		mock_caption_prefix = '1.'
		preknown_caption = False
		prematured_annex = False
		end_of_legal_body = False
		# current_index = 0
		# prev_index = -1
		# last_paragraph_fs = ''
		# wired_document_switch_on = False
		# text = re.sub(r'Article\s\d+}\n\\par\s', '1.', self._contract_doc)
		#--- remove format paragraph
		text = re.sub(r'\\rtf1[\s,\S]*?\\par\s', '', self._contract_doc)
		#--remove footer paragraph
		text = re.sub(r'\\footery\d+[\s,\S]+?\\par\s', '', self._contract_doc)
		#---remove cross-ref: datafields
		# text = re.sub(r'{\\field[\s,\S]+?\\sectd', '', text))
		text = re.sub(r'_Ref\d+', '', text)
		#--- split the text by \\par
		text = re.sub(
			r'\\par\s', self._paragraph_separator_sub, text)
		#--- remove comments 
		text = re.sub(r'\\atrfend[\s,\S]*?}}}', '', text)
		# remove some special text coloring
		text = re.sub(r'\\bkmk.*?}', '', text)
		# remove images 
		text = re.sub(r'\\picscalex[\s,\S]+?}', '', text)
		# remove underscores
		text = re.sub(r'__+', '', text)
		#--- split the text by special string into different pages
		pages = text.split('\\pagebb')
		#--- loop through every paragraph
		current_paragraph_section = 0
		index_topic_map = {}
		try:
			for page in pages:
				page = page.lower()
				paragraph_count = 0
				#--- remove all decorative tags
				page = re.sub(r'\\[a-z,A-Z,0-9,\-,\*,\',\~]*', '', page)
				page = re.sub(r'\\~', '', page)
				#--- remove cross refs 
				page = re.sub(r'_ref\d+', '', page)
				#--- remove table of content bookmarks
				page = re.sub(r'_toc\d*', '', page)
				#--- remove curly brackets
				page = re.sub(r'[{,}]', '', page)
				#--- destroy TOB - example: Cont_0014.rtf
				# Otherwise sometimes >1 DEFINITIONS in the next step will be spotted
				page = re.sub(r'toc[\S,\s]*pageref[\s,\S]*pageref', '', page)
				# TODOs:
				# - hyphen in title
				# - space and line break in title
				# skip definition
				page = page.strip()
				# remove space at front just for non-legal body detection
				pattern = re.compile(
					'^[\\s, \\n, \\t, '+self._paragraph_separator_sub+']+'
				)
				page = re.sub(pattern, '', page)
				paragraphs = page.split(self._paragraph_separator_sub)
				
				for paragraph in paragraphs:
					paragraph_count += 1
					paragraph_raw = paragraph

					# #--- remove all decorative tags
					# paragraph = re.sub(r'\\[a-z,A-Z,0-9,\-,\*,\']*', '', paragraph)
					# paragraph = re.sub(r'\\~', '', paragraph)
					# #--- remove cross refs 

					# #TODO: something left! a long hash string eft like this

					# # 
					# paragraph = re.sub(r'_ref\d+', '', paragraph)
					# #--- remove table of content bookmarks
					# paragraph = re.sub(r'_toc\d*', '', paragraph)
					# #--- remove curly brackets
					# paragraph = re.sub(r'[{,}]', '', paragraph)
					# #--- destroy TOB - example: Cont_0014.rtf
					# # Otherwise sometimes >1 DEFINITIONS in the next step will be spotted
					# paragraph = re.sub(r'toc[\S,\s]*pageref[\s,\S]*pageref', '', paragraph)
					# TODOs:
					# - hyphen in title
					# - space and line break in title

					# remove space at front just for caption detection
					paragraph = re.sub(r'^\s+', '', paragraph)
					
					# if a paragraph starts with Appendix then jump out 
					# and stop looping through the rest of doc as those are not needed
					# 'appendix' 'schedule' 'exhibit' 'annex'
					# trait: page title

					if paragraph_count == 1 and \
						self._non_legal_content_check(paragraph):
						break

					if len(derived_observations) > 0 and \
						self._non_legal_content_check(paragraph):
						# if paragraph starts with these words and
						# body has already been caught 
						# it indicates the end of legal body
						end_of_legal_body = True
						break

					# captions indicated by 'Article' keyword
					# special treatment for Cont_0371.rtf
					if preknown_caption:
						paragraph = mock_caption_prefix + paragraph
					if re.search(r'^article\s\d+', paragraph):
						preknown_caption = True
					else:
						preknown_caption = False

					# replace line break
					paragraph = re.sub(r'\s*\n\s*', '', paragraph)
					# replace consecutive spaces to just one space
					# paragraph = re.sub(r'\s\s+', ' ', paragraph)

					# skip definition
					# paragraph = paragraph.strip()
					if len(paragraph) == 0 or \
						not re.search(r'[a-zA-Z]', paragraph):
						# skip it when
						# - paragraph is empty
						# - only digits and special chars
						continue

					# replace all white space
					no_space_string = re.sub(r'[\s, \n]', '', paragraph)
					# this caption search has excluded nested bullet numbers
					# eg. 1.1.3 a sentence followed 
					# will be excluded
					normal_caption_search = re.search(
						r'^(\d+)[\., \s, \)]*((?![0-9, \., \(, \", \*, \']).+)', no_space_string)

					# last step cleaning
									# denoise text
					paragraph = self.denoise_text(paragraph)
					# replace contractions
					# eg. don't --> do n't
					paragraph = self.replace_contractions(paragraph)

					if (
							normal_caption_search and \
							not re.search(r'\\outlinelevel1', paragraph_raw)
							# not prematured_annex
						):
						# define caption paragraph capture rules
						# - caption-like paragraph
						# - the captured index is increasing OR 
						# it's the first time a caption captured
						# speical treatment for excluding normal clause in definition
						# as caption, eg. Cont_0357.rtf
						# if caption index "1" appears twice then this is the abnormal one
						capcap = re.search(
							r'(\d+)\.?\)?\[?\s*([\w, \s, \-, \,]+)', paragraph
						)

						current_label = capcap.group(2).lower().strip()

						current_label = re.sub(r'\s\s+', ' ', current_label)

						# clear BOW in the case when _multiple_paragraphs is set True
						if self._multiple_paragraphs and \
							same_topic_BOW:
							# if there is something in current BOW, add it first to output
							# as it comes to a finish for previous topic
							derived_observations.append(same_topic_BOW)
							topic_file_pointer.append(
								same_topic_BOW + \
									[
										current_label,
										os.path.basename(self._file_path),
									]
							)
							tfidf_calculation_corpus.append(' '.join(
								[k for k in same_topic_BOW[0]]
							))
							tfidf_calculation_topic.append(int(same_topic_BOW[1]))
							same_topic_BOW = []
						# concatenate broken label - this happen sometimes

						# current_paragraph_section = caption_search.group(1)
						# index_topic_map[]
						# current_label = 'definition'				
					else:
						# no caption-like string is found
						if self._definition_check(current_label) or \
							current_label == '' or \
							not re.search(r'[a-zA-Z]', current_label):
							# Skip when
							# * it is definition
							# * it does not belong to any label			
							# # save font size to last_paragraph_fs for next
							# last_paragraph_fs = font_size

							continue
						else:
							# if no special skipping needed then record it
							# in bag of words
							# remove bullet numbering & letters (roman)
							paragraph = re.sub(r'^[\d, \.]*', '', paragraph)
							paragraph = re.sub(r'[\(]*\w\)', '', paragraph)
							# remove cross ref without links
							paragraph = re.sub(r'[clause, section]\s\d+[\.\d+]*', '', paragraph)
							# replace tilde with space which sometimes happen
							# eg. Cont_0001.rtf section 6.3
							paragraph = re.sub(r'~', ' ', paragraph)
							#TODO: tokenize that easy? white spaces?
							raw_word_vec = self._tokenize(paragraph)
							# normalize the tokens
							raw_word_vec = self.normalize(raw_word_vec)
							# remove tokens with only non-alphabetical chars
							raw_word_vec = [w for w in raw_word_vec if re.search(r'\w', w)]

							paragraph_output = self._word_scan(
								raw_word_vec, current_label, self._strigent_topic
							)
							# if len(paragraph_output) == 0:
							if len(paragraph_output) == 0:
								continue
							else:
								if paragraph_output[0] == {}:
									continue
								
								if self._multiple_paragraphs:
									if same_topic_BOW:
										# combine paragraphs in same topic as one clause
										same_topic_BOW = [update_dict_by_accumulate(
											same_topic_BOW[0],
											paragraph_output[0]
										)] + [paragraph_output[1]]

									else:
										same_topic_BOW = paragraph_output
								else:
									derived_observations.append(
										paragraph_output	
									)
									topic_file_pointer.append(
										paragraph_output + \
											(
												current_label,
												os.path.basename(self._file_path),
											)
									)
									# for tfidf calculation
									tfidf_calculation_corpus.append(' '.join(
										[k for k in paragraph_output[0]]
									))
									tfidf_calculation_topic.append(int(paragraph_output[1]))
				if end_of_legal_body:
					break

			if self._embedding_method == 'tfidf':
				derived_observations = []
				vectorizer = TfidfVectorizer()
				tfidf_mat = vectorizer.fit_transform(
					tfidf_calculation_corpus).toarray()
				words_for_tfidf_mat = vectorizer.get_feature_names()
				# pd.DataFrame(tfidf_mat)
				for i in range(tfidf_mat.shape[0]):
					derived_observations.append([
						dict(zip(words_for_tfidf_mat, tfidf_mat[i])), 
						tfidf_calculation_topic[i]
					])

		except Exception as e:
			logger.error(self._file_path+"\n Some error happen in preprocessing")

		if len(derived_observations) == 0:
			logger.error(self._file_path+"\n Empty")

		return {
			'derived_observations': derived_observations,
			'topic_file_pointer': topic_file_pointer
		}
