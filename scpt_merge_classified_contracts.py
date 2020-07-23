# To classify all contracts samples and generate an assembly document
# Author: Jintao Long <jintaolong@brookes.ac.uk>

__doc__ = """
"""
import re
import os
import pickle
import datetime 

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from os import listdir
from collections import OrderedDict
from typing import List

from preprocessing import RtfPreprocessing

CONTRACT_PATH = 'contracts/'
model_file = open('models/20200318150525.sav', 'rb')
MODEL = pickle.load(model_file)
ANSWERS_META = {
    'correct': 0,
    'close': 0,
    'miss': 0
}
LABEL_VALUE_MAP = {
    2: "commencement/duration",
    3: "tupe",
    4: "supplier responsibilities (may contain warranties)",
    5: "customer obligations (may contain warranties)",
    6: "non solicitation",
    7: "change control",
    8: "charges, payments and payment terms",
    9: "prices",
    10: "price reviews (cpi/rpi/annual)", #""
    11: "expenses",
    12: "ipr",
    13: "compliance with laws and policies",
    14: "data protection",
    15: "confidentiality",
    16: "limitation of liability",
    #TODO: liability and indemnity ---> what this belongs to?
    17: "termination",
    18: "consequences of termination",
    19: "force majeure",
    20: "costs",
    21: "assignment",
    22: "variation",
    23: "waiver",
    24: "rights and remedies",
    25: "further assurance",
    26: "severance",
    27: "entire agreement",
    28: "conflict",
    29: "no partnership or agency",
    30: "third party rights",
    31: "several liability",
    32: "notices",
    33: "counterparts",
    34: "adr",
    35: "expert determination",
    36: "warranties",
    37: "indemnities",
    38: "insurance",
    39: "licence grant (software, images, etc)",
    40: "scope (may be in separate schedule, or could be in body of agreement)",
    41: "services",
    42: "deliverables",
    43: "testing (pre-delivery and acceptance)",
    44: "acceptance",
    45: "ownership",
    46: "rights in the software (transfer, reproduction, use, adaptation)",
    47: "escrow (source code)",
    48: "support (may be in separate schedule)",
    49: "training",
    50: "personnel",
    51: "anti bribery",
    52: "export control",
    53: "insurance",
    54: "subscriptions",
    55: "supply of products",
    56: "forecasts",
    57: "orders",
    58: "quality/packing",
    59: "delivery",
    60: "acceptance",
    61: "faulty products",
    62: "returns procedure",
    63: "title and risk",
    64: "recalls"
}

class ContractFullScanner(RtfPreprocessing):
    def doc_handler_setter(self, doc):
        self._doc_handler = doc

    def _numerize_texts(self) -> List:
        self._doc_handler.add_heading(self._file_path, level=1)
        derived_observations = []
        same_topic_BOW = []
        # output_
        # add a file to point topic to file
        # only used as reference for research
        topic_file_pointer = []
        tfidf_calculation_corpus = []
        tfidf_calculation_topic = []
        current_label = ''
        mock_caption_prefix = '1.'
        preknown_caption = False
        prematured_annex = False
        end_of_legal_body = False
        # current_index = 0
        # prev_index = -1
        # last_paragraph_fs = ''
        # wired_document_switch_on = False
        # text = re.sub(r'Article\s\d+}\n\\par\s', '1.', self._contract_doc)
        #--- remove format paragraph
        text = self._full_text_pre_clean()
        #--- split the text by special string into different pages
        pages = text.split('\\pagebb')
        #--- loop through every paragraph
        current_paragraph_section = 0
        index_topic_map = {}
        # try:
        for page in pages:
            paragraph_count = 0
            page = self._full_page_pre_clean(page)
            paragraphs = page.split(self._paragraph_separator_sub)
            
            for paragraph in paragraphs:
                paragraph_count += 1
                paragraph_raw = paragraph
                # remove space at front just for caption detection
                paragraph = re.sub(r'^\s+', '', paragraph)
                
                # if a paragraph starts with Appendix then jump out 
                # and stop looping through the rest of doc as those are not needed
                # 'appendix' 'schedule' 'exhibit' 'annex'
                # trait: page title

                if paragraph_count == 1 and \
                    self._non_legal_content_check(paragraph):
                    break

                if len(derived_observations) > 0 and \
                    self._non_legal_content_check(paragraph):
                    # if paragraph starts with these words and
                    # body has already been caught 
                    # it indicates the end of legal body
                    end_of_legal_body = True
                    break

                # captions indicated by 'Article' keyword
                # special treatment for Cont_0371.rtf
                if preknown_caption:
                    paragraph = mock_caption_prefix + paragraph
                if re.search(r'^article\s\d+', paragraph):
                    preknown_caption = True
                else:
                    preknown_caption = False

                # replace line break
                paragraph = re.sub(r'\s*\n\s*', '', paragraph)
                # replace consecutive spaces to just one space
                # paragraph = re.sub(r'\s\s+', ' ', paragraph)

                # skip definition
                # paragraph = paragraph.strip()
                if len(paragraph) == 0 or \
                    not re.search(r'[a-zA-Z]', paragraph):
                    # skip it when
                    # - paragraph is empty
                    # - only digits and special chars
                    continue

                # replace all white space
                no_space_string = re.sub(r'[\s, \n]', '', paragraph)
                # this caption search has excluded nested bullet numbers
                # eg. 1.1.3 a sentence followed 
                # will be excluded
                normal_caption_search = re.search(
                    r'^(\d+)[\., \s, \)]*((?![0-9, \., \(, \", \*, \']).+)', no_space_string)

                # last step cleaning
                                # denoise text
                paragraph = self.denoise_text(paragraph)
                # replace contractions
                # eg. don't --> do n't
                paragraph = self.replace_contractions(paragraph)

                if (
                        normal_caption_search and \
                        not re.search(r'\\outlinelevel1', paragraph_raw)
                        # not prematured_annex
                    ):
                    # define caption paragraph capture rules
                    # - caption-like paragraph
                    # - the captured index is increasing OR 
                    # it's the first time a caption captured
                    # speical treatment for excluding normal clause in definition
                    # as caption, eg. Cont_0357.rtf
                    # if caption index "1" appears twice then this is the abnormal one
                    capcap = re.search(
                        r'(\d+)\.?\)?\[?\s*([\w, \s, \-, \,]+)', paragraph
                    )
                    try:
                        current_label = capcap.group(2).lower().strip()
                    except Exception as e:
                        pass

                    current_label = re.sub(r'\s\s+', ' ', current_label)
                    
                    par = self._doc_handler.add_paragraph()
                    run = par.add_run(current_label)
                    run.bold = True
                else:
                    # no caption-like string is found
                    if self._definition_check(current_label) or \
                        current_label == '' or \
                        not re.search(r'[a-zA-Z]', current_label):
                        # Skip when
                        # * it is definition
                        # * it does not belong to any label         
                        # # save font size to last_paragraph_fs for next
                        # last_paragraph_fs = font_size

                        continue
                    else:
                        # if no special skipping needed then record it
                        # in bag of words
                        # remove bullet numbering & letters (roman)
                        paragraph = re.sub(r'^[\d, \.]*', '', paragraph)
                        paragraph = re.sub(r'[\(]*\w\)', '', paragraph)
                        # remove cross ref without links
                        paragraph = re.sub(r'[clause, section]\s\d+[\.\d+]*', '', paragraph)
                        # replace tilde with space which sometimes happen
                        # eg. Cont_0001.rtf section 6.3
                        paragraph = re.sub(r'~', ' ', paragraph)
                        #TODO: tokenize that easy? white spaces?
                        raw_word_vec = self._tokenize(paragraph)
                        # normalize the tokens
                        raw_word_vec = self.normalize(raw_word_vec)
                        # remove tokens with only non-alphabetical chars
                        raw_word_vec = [w for w in raw_word_vec if re.search(r'\w', w)]
                        
                        if len(raw_word_vec) <= 13:
                            continue
                        
                        paragraph_output = self._word_scan(
                            raw_word_vec, current_label, self._strigent_topic
                        )
                        # if len(paragraph_output) == 0:
                        if len(paragraph_output) == 0:
                            continue
                        else:
                            if paragraph_output[0] == {}:
                                continue
                            
                            derived_observations.append(
                                paragraph_output    
                            )
                            topic_file_pointer.append(
                                paragraph_output + \
                                    (
                                        current_label,
                                        os.path.basename(self._file_path),
                                    )
                            )

                            # Predicting stuff
                            bow = paragraph_output[0]
                            predictions = MODEL.prob_classify(bow)
                            probs = OrderedDict()
                            probs_ind = OrderedDict()
                            for topic in LABEL_VALUE_MAP:
                                probs[LABEL_VALUE_MAP[topic]] = predictions.prob(topic)
                                probs_ind[topic] = predictions.prob(topic)
                            
                            probs = OrderedDict(
                                sorted(probs.items(), key=lambda x: x[1], reverse=True)
                            )
                            probs_ind = OrderedDict(
                                sorted(probs_ind.items(), key=lambda x: x[1], reverse=True)
                            )
                            top_5 = OrderedDict(list(probs.items())[0:5])
                            top_5_ind = OrderedDict(list(probs_ind.items())[0:5])
                            probs = {x[0]:"{:.2%}".format(x[1]) for x in top_5.items()}
                            pred = MODEL.classify(bow)
                            
                            if int(pred) != int(paragraph_output[1]):
                                # mark it when incorrect
                                par = self._doc_handler.add_paragraph()
                                font = par.add_run(paragraph).font
                                if paragraph_output[1] in top_5_ind.keys():
                                    font.highlight_color = WD_COLOR_INDEX.YELLOW
                                    ANSWERS_META['close'] += 1
                                else:
                                    font.highlight_color = WD_COLOR_INDEX.RED
                                    ANSWERS_META['miss'] += 1
                                # add probability 
                                self._doc_handler.add_paragraph(str(probs))
                            else:
                                # stay normal when it's correct
                                par = self._doc_handler.add_paragraph(paragraph)
                                self._doc_handler.add_paragraph(str(probs))
                                ANSWERS_META['correct'] += 1

            if end_of_legal_body:
                break
        return {
            'derived_observations': derived_observations,
            'topic_file_pointer': topic_file_pointer
        }


if __name__ == '__main__':
    # f = open('full_scan.docx', 'wb')
    # file_path = 'contracts/Cont_0085.rtf'
    file_path = None
    doc = Document()
    contracts_path = os.path.join(
        os.getcwd(),
        CONTRACT_PATH
    )
    file_paths = listdir(contracts_path)
    if file_path:
        file_paths = [
            os.path.join(
                os.getcwd(),
                file_path
            )
        ]
    else:
        file_paths = [os.path.join(contracts_path, p) for p in file_paths]

    for fp in file_paths:
        if fp.endswith('rtf'):
            cfs = ContractFullScanner(
                fp, 
                'count_occurence',
                'heading_as_label',
                True,
                False,
                True,
                4,
                True,
                True
            )
            cfs.doc_handler_setter(doc)
            cfs.bag_of_word_dict_transformer()
        print('File '+fp+' processed')
    doc.save('full_scan.docx')
    print('Correct: {} \n\r Close: {} \n\r Miss: {} \n\r'.format(
        str(ANSWERS_META['correct']),
        str(ANSWERS_META['close']),
        str(ANSWERS_META['miss'])
    ))